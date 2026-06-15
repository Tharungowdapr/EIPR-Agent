import io
import json
import re
from html import escape
from datetime import datetime


def _md_to_html(text: str) -> str:
    if not text:
        return ""
    import markdown as md_lib
    return md_lib.markdown(str(text), output_format="html5")


def _strip_markdown(text: str) -> str:
    if not text:
        return ""
    t = str(text)
    t = re.sub(r'(\*{1,2}|_{1,2})(.+?)\1', r'\2', t)
    t = re.sub(r'^#{1,6}\s+', '', t, flags=re.MULTILINE)
    t = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t)
    t = re.sub(r'^[-*+]\s+', '', t, flags=re.MULTILINE)
    t = re.sub(r'^\d+\.\s+', '', t, flags=re.MULTILINE)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    t = re.sub(r'^>\s+', '', t, flags=re.MULTILINE)
    t = re.sub(r'~~~[\s\S]*?~~~', '', t)
    t = re.sub(r'```[\s\S]*?```', '', t)
    return t


def generate_docx(report: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return _fallback_docx(report)

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    title = report.get("title", "EIPR Case Study Report")
    doc.add_heading(title, level=0)

    abstract = report.get("abstract", "")
    if abstract:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(_strip_markdown(abstract))

    case_study = report.get("case_study", {})
    sections = [
        ("1. Introduction", "introduction"),
        ("2. Opportunity Analysis", "opportunity_analysis"),
        ("3. Business Strategy", "business_strategy"),
        ("4. IP Strategy", "ip_strategy"),
        ("5. Conclusion", "conclusion"),
    ]

    for heading, key in sections:
        content = case_study.get(key, "")
        if content:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(_strip_markdown(content))

    eipr_mapping = report.get("eipr_mapping", [])
    if eipr_mapping:
        doc.add_heading("EIPR Curriculum Mapping", level=1)
        for mapping in eipr_mapping:
            doc.add_heading(f"Unit {mapping.get('unit', '')}: {mapping.get('topic', '')}", level=2)
            doc.add_paragraph(_strip_markdown(mapping.get("coverage", "")))

    learning = report.get("learning_outcomes", [])
    if learning:
        doc.add_heading("Learning Outcomes", level=1)
        for item in learning:
            doc.add_paragraph(_strip_markdown(item), style='List Bullet')

    questions = report.get("discussion_questions", [])
    if questions:
        doc.add_heading("Discussion Questions", level=1)
        for q in questions:
            doc.add_paragraph(_strip_markdown(q), style='List Number')

    keywords = report.get("keywords", [])
    if keywords:
        doc.add_heading("Keywords", level=1)
        doc.add_paragraph(", ".join(keywords))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_pdf_html(report: dict) -> str:
    def esc(v: str) -> str:
        return escape(str(v))

    title = report.get("title", "EIPR Case Study Report")
    abstract = report.get("abstract", "")
    case_study = report.get("case_study", {})
    eipr_mapping = report.get("eipr_mapping", [])
    learning = report.get("learning_outcomes", [])
    questions = report.get("discussion_questions", [])
    keywords = report.get("keywords", [])

    sections_html = ""
    for key, label in [("introduction", "1. Introduction"), ("opportunity_analysis", "2. Opportunity Analysis"),
                       ("business_strategy", "3. Business Strategy"), ("ip_strategy", "4. IP Strategy"),
                       ("conclusion", "5. Conclusion")]:
        content = case_study.get(key, "")
        if content:
            sections_html += f"<h2>{esc(label)}</h2>{_md_to_html(content)}"

    mapping_html = ""
    for m in eipr_mapping:
        mapping_html += f"<h3>Unit {esc(str(m.get('unit', '')))}: {esc(str(m.get('topic', '')))}</h3>{_md_to_html(m.get('coverage', ''))}"

    learning_html = "<ul>" + "".join(f"<li>{esc(l)}</li>" for l in learning) + "</ul>" if learning else ""
    questions_html = "<ol>" + "".join(f"<li>{esc(q)}</li>" for q in questions) + "</ol>" if questions else ""

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{esc(title)}</title>
<style>
body {{ font-family: 'Times New Roman', serif; margin: 1in; font-size: 11pt; line-height: 1.6; }}
h1 {{ font-size: 18pt; text-align: center; margin-bottom: 20px; }}
h2 {{ font-size: 14pt; margin-top: 20px; border-bottom: 1px solid #ccc; }}
h3 {{ font-size: 12pt; margin-top: 15px; }}
p {{ text-align: justify; }}
ul, ol {{ margin-left: 20px; }}
</style></head><body>
<h1>{esc(title)}</h1>
<h2>Abstract</h2>{_md_to_html(abstract)}
{sections_html}
<h2>EIPR Curriculum Mapping</h2>{mapping_html}
<h2>Learning Outcomes</h2>{learning_html}
<h2>Discussion Questions</h2>{questions_html}
<p><strong>Keywords:</strong> {esc(', '.join(keywords))}</p>
</body></html>"""


def _md_inline(text: str) -> str:
    escaped = escape(str(text))
    escaped = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', escaped)
    escaped = re.sub(r'__([^_]+)__', r'<strong>\1</strong>', escaped)
    escaped = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<em>\1</em>', escaped)
    escaped = re.sub(r'(?<!_)_([^_]+)_(?!_)', r'<em>\1</em>', escaped)
    escaped = re.sub(r'`([^`]+)`', r'<code>\1</code>', escaped)
    return escaped


def _render_value(val, indent: int = 0) -> str:
    prefix = "  " * indent
    if val is None:
        return ""
    if isinstance(val, dict):
        parts = []
        for k, v in val.items():
            label = k.replace("_", " ").title()
            rendered = _render_value(v, indent + 1)
            if rendered:
                if isinstance(v, (dict, list)):
                    parts.append(f"{prefix}<h4 style='font-size:10pt; margin:8px 0 4px;'>{escape(label)}</h4>{rendered}")
                else:
                    parts.append(f"{prefix}<p style='margin:2px 0;'><strong>{escape(label)}:</strong> {rendered}</p>")
        return "\n".join(parts)
    if isinstance(val, list):
        if not val:
            return ""
        if all(isinstance(v, str) for v in val):
            return "<ul>" + "".join(f"<li>{_md_inline(v)}</li>" for v in val) + "</ul>"
        if all(isinstance(v, dict) for v in val):
            return "\n".join(_render_value(v, indent) + "<hr style='border:none; border-top:1px solid #eee; margin:6px 0;'/>" for v in val)
        return "<ul>" + "".join(f"<li>{_md_inline(str(v))}</li>" for v in val) + "</ul>"
    if isinstance(val, bool):
        return "Yes" if val else "No"
    try:
        v = float(val)
        if v >= 1e7:
            return f"₹{v/1e7:.2f} Cr"
        elif v >= 1e5:
            return f"₹{v/1e5:.2f} L"
        if val == int(val):
            return f"₹{int(v):,}" if v >= 0 else str(val)
    except (ValueError, TypeError):
        pass
    return _md_inline(str(val))


def generate_full_project_pdf_html(project: dict, outputs: dict) -> str:
    def esc(v: str) -> str:
        return escape(str(v))

    title = project.get("title", "EIPR Project Report")
    domain = project.get("domain", "")
    description = project.get("input_text", "")
    context = project.get("user_context", "")

    sections = [f"<h1>{esc(title)}</h1>"]
    sections.append(f'<p><strong>Domain:</strong> {esc(domain)} | <strong>Generated:</strong> {datetime.now().strftime("%B %d, %Y")}</p>')
    sections.append('<hr/>')
    sections.append(f'<h2>Project Description</h2>{_md_to_html(description)}')
    if context:
        sections.append(f'<h3>User Context</h3>{_md_to_html(context)}')

    opp_data = outputs.get("opportunities") or {}
    opportunities = opp_data.get("opportunities", [])

    if opp_data:
        sections.append("<h2>1. Domain Analysis</h2>")
        for key in ["summary", "current_market_landscape", "user_gaps", "india_relevance"]:
            val = opp_data.get(key)
            if val:
                label = key.replace("_", " ").title()
                sections.append(f"<h3>{esc(label)}</h3>{_md_to_html(str(val))}")
        for key in ["emerging_trends", "impact_factors"]:
            val = opp_data.get(key)
            if val and isinstance(val, list):
                label = key.replace("_", " ").title()
                sections.append(f"<h3>{esc(label)}</h3><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")

    if opportunities:
        sections.append("<h2>2. Opportunity Analysis</h2>")
        for i, opp in enumerate(opportunities, 1):
            sections.append(f"<h3>Opportunity {i}: {esc(opp.get('title', ''))}</h3>")
            sections.append(_md_to_html(opp.get('description', '')))
            metric_rows = [
                ["Feasibility Score", f"{opp.get('feasibility_score', 'N/A')}/10"],
                ["Innovation Level", esc(opp.get('innovation_level', 'N/A'))],
                ["Market Gap", esc(opp.get('market_gap', 'N/A'))],
                ["Target Customer", esc(opp.get('target_customer', 'N/A'))],
                ["TAM", esc(opp.get('tam', 'N/A'))],
                ["SAM", esc(opp.get('sam', 'N/A'))],
                ["SOM", esc(opp.get('som', 'N/A'))],
                ["IP Potential", esc(opp.get('ip_potential', 'N/A'))],
                ["Government Alignment", esc(opp.get('government_alignment', 'N/A'))],
                ["Regulatory Notes", esc(opp.get('regulatory_notes', 'N/A'))],
                ["India Specific Risk", esc(opp.get('india_specific_risk', 'N/A'))],
            ]
            sections.append(f"<h4>Key Metrics</h4>{_table(['Metric', 'Value'], metric_rows)}")
            for field, flabel in [("entrepreneur_type", "Entrepreneur Type"), ("entrepreneur_type_reasoning", "Entrepreneur Type Reasoning")]:
                val = opp.get(field)
                if val:
                    sections.append(f"<h4>{esc(flabel)}</h4>{_md_to_html(str(val))}")
            for field, flabel in [("success_factors", "Success Factors"), ("challenges", "Challenges")]:
                val = opp.get(field)
                if val and isinstance(val, list):
                    sections.append(f"<h4>{esc(flabel)}</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")
            if i < len(opportunities):
                sections.append("<hr/>")

        ei = opp_data.get("entrepreneurial_insights", {})
        if ei:
            sections.append("<h2>Entrepreneurial Insights</h2>")
            for key in ["traits_needed", "myths_busted"]:
                val = ei.get(key)
                if val and isinstance(val, list):
                    label = key.replace("_", " ").title()
                    sections.append(f"<h3>{esc(label)}</h3><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")
            for key in ["team_requirements", "recommended_approach"]:
                val = ei.get(key)
                if val:
                    label = key.replace("_", " ").title()
                    sections.append(f"<h3>{esc(label)}</h3>{_md_to_html(str(val))}")

    for idx in range(5):
        ip = outputs.get(f"ip_analysis_{idx}")
        if ip:
            sections.append(f"<h2>3. IP Analysis (Opportunity {idx + 1})</h2>")
            for key, flabel in [("analysis", "Analysis"), ("summary", "Summary")]:
                val = ip.get(key)
                if val:
                    sections.append(f"<h3>{esc(flabel)}</h3>{_render_value(val)}")

            pa = ip.get("patent_analysis") or {}
            if pa:
                sections.append("<h3>Patent Analysis (Patents Act 1970)</h3>")
                for key in ["patentability_assessment", "prior_art_search_strategy", "freedom_to_operate", "section_3_considerations"]:
                    val = pa.get(key)
                    if val:
                        label = key.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4>{_md_to_html(str(val))}")
                val = pa.get("likely_cpc_codes") or pa.get("indian_ipc_classes")
                if val and isinstance(val, list):
                    sections.append(f"<h4>Patent Classification Codes</h4><p>{esc(', '.join(val))}</p>")
                score = pa.get("estimated_patentability_score")
                if score:
                    sections.append(f"<p><strong>Estimated Patentability Score:</strong> {esc(str(score))}/10</p>")
                fs = pa.get("filing_strategy") or {}
                if fs:
                    sections.append("<h4>Filing Strategy</h4>")
                    fs_rows = []
                    for fsk in ["type", "jurisdictions", "timeline", "startup_fee_discount", "estimated_cost_inr"]:
                        fsv = fs.get(fsk)
                        if fsv:
                            if isinstance(fsv, list):
                                fsv = ", ".join(fsv)
                            fs_rows.append([fsk.replace("_", " ").title(), esc(str(fsv))])
                    if fs_rows:
                        sections.append(_table(["Parameter", "Value"], fs_rows))

            ta = ip.get("trademark_analysis") or {}
            if ta:
                sections.append("<h3>Trademark Analysis (Trade Marks Act 1999)</h3>")
                for key in ["registration_strategy", "class_description"]:
                    val = ta.get(key)
                    if val:
                        label = key.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4>{_md_to_html(str(val))}")
                for key in ["protectable_elements", "nice_classes"]:
                    val = ta.get(key)
                    if val and isinstance(val, list):
                        label = key.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4><p>{esc(', '.join(val))}</p>")
                if ta.get("clearance_required") is not None:
                    sections.append(f"<p><strong>Clearance Required:</strong> {'Yes' if ta['clearance_required'] else 'No'}</p>")

            ca = ip.get("copyright_analysis") or {}
            if ca:
                sections.append("<h3>Copyright Analysis (Copyright Act 1957)</h3>")
                for key in ["ownership_notes", "licensing_strategy"]:
                    val = ca.get(key)
                    if val:
                        label = key.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4>{_md_to_html(str(val))}")
                val = ca.get("copyrightable_elements")
                if val and isinstance(val, list):
                    sections.append(f"<h4>Copyrightable Elements</h4><p>{esc(', '.join(val))}</p>")

            tsa = ip.get("trade_secret_analysis") or {}
            if tsa:
                sections.append("<h3>Trade Secret Analysis</h3>")
                for key in ["recommendation"]:
                    val = tsa.get(key)
                    if val:
                        sections.append(f"<h4>{esc(key.replace('_', ' ').title())}</h4>{_md_to_html(str(val))}")
                for key in ["trade_secret_candidates", "protection_measures"]:
                    val = tsa.get(key)
                    if val and isinstance(val, list):
                        label = key.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")

            ips = ip.get("ip_strategy") or {}
            if ips:
                sections.append("<h3>IP Strategy Recommendations</h3>")
                for key in ["portfolio_roadmap", "commercialization_path", "enforcement_strategy"]:
                    val = ips.get(key)
                    if val:
                        label = key.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4>{_md_to_html(str(val))}")
                for k in ["budget_estimate", "estimated_ip_value"]:
                    val = ips.get(k)
                    if val:
                        sections.append(f"<p><strong>{k.replace('_', ' ').title()}:</strong> {esc(str(val))}</p>")
                val = ips.get("timeline_months")
                if val:
                    sections.append(f"<p><strong>Timeline:</strong> {esc(str(val))} months</p>")
                val = ips.get("key_recommendations")
                if val and isinstance(val, list):
                    sections.append("<h4>Key Recommendations</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")

    for idx in range(5):
        bp = outputs.get(f"business_plan_{idx}")
        if bp:
            sections.append(f"<h2>4. Business Strategy (Opportunity {idx + 1})</h2>")
            for key in ["executive_summary", "company_description"]:
                val = bp.get(key)
                if val:
                    label = key.replace("_", " ").title()
                    sections.append(f"<h3>{esc(label)}</h3>{_md_to_html(str(val))}")

            ma = bp.get("market_analysis") or {}
            if ma:
                sections.append("<h3>Market Analysis</h3>")
                for k in ["tam", "sam", "som", "competitor_landscape"]:
                    val = ma.get(k)
                    if val:
                        sections.append(f"<p><strong>{k.upper()}:</strong> {esc(str(val))}</p>")
                val = ma.get("market_trends")
                if val and isinstance(val, list):
                    sections.append("<h4>Market Trends</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")

            bm = bp.get("business_model") or {}
            if bm:
                sections.append("<h3>Business Model</h3>")
                for k in ["pricing_strategy", "unit_economics"]:
                    val = bm.get(k)
                    if val:
                        label = k.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4>{_md_to_html(str(val))}")
                val = bm.get("revenue_streams")
                if val and isinstance(val, list):
                    sections.append("<h4>Revenue Streams</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")

            sa = bp.get("strategic_analysis") or {}
            if sa:
                sections.append("<h3>Strategic Analysis</h3>")
                swoc = sa.get("swoc") or {}
                for k in ["strengths", "weaknesses", "opportunities", "challenges"]:
                    val = swoc.get(k)
                    if val and isinstance(val, list):
                        label = k.replace("_", " ").title()
                        sections.append(f"<h4>SWOC — {esc(label)}</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")
                vmgo = sa.get("vmgo") or {}
                for k in ["vision", "mission"]:
                    val = vmgo.get(k)
                    if val:
                        sections.append(f"<h4>VMGO — {k.title()}</h4>{_md_to_html(str(val))}")
                val = vmgo.get("goals") or vmgo.get("objectives")
                if val and isinstance(val, list):
                    label = "Goals & Objectives"
                    sections.append(f"<h4>{label}</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")
                ps = sa.get("porter_strategy") or {}
                if ps:
                    sections.append("<h4>Porter's Strategy</h4>")
                    for k in ["recommended", "reasoning"]:
                        val = ps.get(k)
                        if val:
                            sections.append(f"<p>{k.title()}: {esc(str(val))}</p>")
                val = sa.get("competitive_advantage")
                if val:
                    sections.append(f"<h4>Competitive Advantage</h4>{_md_to_html(str(val))}")

            mp = bp.get("marketing_plan") or {}
            if mp:
                sections.append("<h3>Marketing Plan</h3>")
                stp = mp.get("stp") or {}
                for k in ["segments", "target", "positioning"]:
                    val = stp.get(k)
                    if val:
                        if isinstance(val, list):
                            sections.append(f"<h4>STP — {k.title()}</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")
                        else:
                            sections.append(f"<h4>STP — {k.title()}</h4>{_md_to_html(str(val))}")
                val = mp.get("uvp")
                if val:
                    sections.append(f"<h4>Unique Value Proposition</h4>{_md_to_html(str(val))}")
                fps = mp.get("four_ps") or {}
                for k in ["product", "price", "place", "promotion"]:
                    val = fps.get(k)
                    if val:
                        sections.append(f"<h4>4Ps — {k.title()}</h4>{_md_to_html(str(val))}")
                dm = mp.get("digital_marketing") or {}
                for k in ["seo_strategy", "sem_strategy", "social_media", "content_marketing"]:
                    val = dm.get(k)
                    if val:
                        label = k.replace("_", " ").title()
                        sections.append(f"<h4>Digital Marketing — {esc(label)}</h4>{_md_to_html(str(val))}")

            gs = bp.get("growth_strategy") or {}
            if gs:
                sections.append("<h3>Growth Strategy</h3>")
                for k in ["organic_growth", "scaling_roadmap"]:
                    val = gs.get(k)
                    if val:
                        label = k.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4>{_md_to_html(str(val))}")
                for k in ["strategic_alliances", "exit_strategies"]:
                    val = gs.get(k)
                    if val and isinstance(val, list):
                        label = k.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")

            cs = bp.get("cost_structure")
            if cs and isinstance(cs, list):
                rows = [[esc(item.get("category", "")), _render_value(item.get("amount", 0)), esc(item.get("notes", ""))] for item in cs if isinstance(item, dict)]
                if rows:
                    sections.append("<h3>Cost Structure</h3>")
                    sections.append(_table(["Category", "Amount", "Notes"], rows))

    for idx in range(5):
        fin = outputs.get(f"financial_{idx}")
        if fin:
            sections.append(f"<h2>5. Financial Analysis (Opportunity {idx + 1})</h2>")

            sc = fin.get("startup_costs") or {}
            if sc:
                sections.append("<h3>Startup Costs</h3>")
                cost_rows = []
                for k in ["development", "marketing", "operations", "legal_ip", "office_infrastructure"]:
                    val = sc.get(k)
                    if val:
                        cost_rows.append([k.replace("_", " ").title(), _render_value(val)])
                for k in ["total_initial_investment", "monthly_burn_rate"]:
                    val = sc.get(k)
                    if val:
                        cost_rows.append([k.replace("_", " ").title(), _render_value(val)])
                if cost_rows:
                    sections.append(_table(["Category", "Amount"], cost_rows))
                cb = sc.get("cost_breakdown")
                if cb and isinstance(cb, list):
                    sub_rows = [[esc(item.get("category", "")), _render_value(item.get("amount", 0)), esc(item.get("notes", ""))] for item in cb if isinstance(item, dict)]
                    if sub_rows:
                        sections.append("<h4>Detailed Cost Breakdown</h4>")
                        sections.append(_table(["Category", "Amount", "Notes"], sub_rows))

            rp = fin.get("revenue_projections") or {}
            if rp:
                sections.append("<h3>Revenue Projections</h3>")
                for year_key in ["year_1", "year_2", "year_3"]:
                    yr = rp.get(year_key) or {}
                    if yr:
                        label = year_key.replace("_", " ").title()
                        yr_rows = []
                        for k in ["mrr", "arr"]:
                            val = yr.get(k)
                            if val:
                                yr_rows.append([k.upper(), _render_value(val)])
                        if yr_rows:
                            sections.append(f"<h4>{esc(label)}</h4>{_table(['Metric', 'Value'], yr_rows)}")
                        qb = yr.get("quarterly_breakdown")
                        if qb and isinstance(qb, list):
                            qrows = [[esc(item.get("q", "")), _render_value(item.get("revenue", 0))] for item in qb if isinstance(item, dict)]
                            if qrows:
                                sections.append(f"<h4>{esc(label)} — Quarterly Breakdown</h4>{_table(['Quarter', 'Revenue'], qrows)}")
                val = rp.get("key_assumptions")
                if val and isinstance(val, list):
                    sections.append("<h4>Key Assumptions</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")
                val = rp.get("growth_rate_pct")
                if val:
                    sections.append(f"<p><strong>Growth Rate:</strong> {esc(str(val))}%</p>")

            be = fin.get("break_even_analysis") or {}
            if be:
                sections.append("<h3>Break-Even Analysis</h3>")
                be_rows = []
                for k in ["fixed_costs_monthly", "variable_costs_per_unit", "unit_price", "contribution_margin", "break_even_revenue"]:
                    val = be.get(k)
                    if val:
                        be_rows.append([k.replace("_", " ").title(), _render_value(val)])
                for k in ["break_even_units_monthly", "break_even_months"]:
                    val = be.get(k)
                    if val:
                        be_rows.append([k.replace("_", " ").title(), esc(str(val))])
                if be_rows:
                    sections.append(_table(["Metric", "Value"], be_rows))

            fs = fin.get("funding_strategy") or {}
            if fs:
                sections.append("<h3>Funding Strategy</h3>")
                val = fs.get("total_funding_needed")
                if val:
                    sections.append(f"<p><strong>Total Funding Needed:</strong> {_render_value(val)}</p>")
                rm = fs.get("recommended_mix") or {}
                if rm:
                    mix_rows = [[k.replace("_", " ").title(), esc(str(v))] for k, v in rm.items() if v]
                    if mix_rows:
                        sections.append("<h4>Recommended Funding Mix</h4>")
                        sections.append(_table(["Source", "Allocation"], mix_rows))
                val = fs.get("funding_roadmap")
                if val:
                    sections.append(f"<h4>Funding Roadmap</h4>{_md_to_html(str(val))}")
                val = fs.get("investor_readiness_score")
                if val:
                    sections.append(f"<p><strong>Investor Readiness Score:</strong> {esc(str(val))}/10</p>")
                for k in ["indian_grant_eligibility", "recommended_investors"]:
                    val = fs.get(k)
                    if val and isinstance(val, list):
                        label = k.replace("_", " ").title()
                        sections.append(f"<h4>{esc(label)}</h4><ul>" + "".join(f"<li>{esc(v)}</li>" for v in val) + "</ul>")

            fm = fin.get("financial_metrics") or {}
            if fm:
                sections.append("<h3>Financial Metrics</h3>")
                fm_rows = []
                for k in ["gross_margin_pct", "net_margin_year_1", "net_margin_year_2", "net_margin_year_3"]:
                    val = fm.get(k)
                    if val:
                        fm_rows.append([k.replace("_", " ").title(), f"{esc(str(val))}%"])
                for k in ["cac", "ltv"]:
                    val = fm.get(k)
                    if val:
                        fm_rows.append([k.upper(), _render_value(val)])
                for k in ["ltv_cac_ratio", "payback_period_months", "roi_year_3_pct"]:
                    val = fm.get(k)
                    if val:
                        fm_rows.append([k.replace("_", " ").title(), esc(str(val))])
                val = fm.get("tax_considerations")
                if val:
                    fm_rows.append(["Tax Considerations", esc(str(val))])
                if fm_rows:
                    sections.append(_table(["Metric", "Value"], fm_rows))

    report = outputs.get("report", {})
    if report:
        sections.append("<h2>6. Final Report</h2>")
        val = report.get("abstract")
        if val:
            sections.append(f"<h3>Abstract</h3>{_md_to_html(val)}")

        case_study = report.get("case_study", {})
        for key, label in [("introduction", "Introduction"), ("opportunity_analysis", "Opportunity Analysis"),
                           ("business_strategy", "Business Strategy"), ("ip_strategy", "IP Strategy"),
                           ("conclusion", "Conclusion")]:
            val = case_study.get(key)
            if val:
                sections.append(f"<h3>{esc(label)}</h3>{_md_to_html(val)}")

        em = report.get("eipr_mapping", [])
        if em:
            sections.append("<h3>EIPR Curriculum Mapping</h3>")
            for m in em:
                unit = esc(str(m.get("unit", "")))
                topic = esc(str(m.get("topic", "")))
                coverage = _md_to_html(str(m.get("coverage", "")))
                sections.append(f"<h4>Unit {unit}: {topic}</h4>{coverage}")
                for fk, fl in [("indian_examples", "Indian Examples"), ("indian_law_references", "Indian Law References")]:
                    fv = m.get(fk)
                    if fv and isinstance(fv, list):
                        sections.append(f"<p><strong>{fl}:</strong></p><ul>" + "".join(f"<li>{esc(v)}</li>" for v in fv) + "</ul>")

        for field, flabel in [("learning_outcomes", "Learning Outcomes"), ("key_takeaways", "Key Takeaways"),
                               ("discussion_questions", "Discussion Questions")]:
            val = report.get(field)
            if val and isinstance(val, list):
                tag = "ol" if field == "discussion_questions" else "ul"
                sections.append(f"<h3>{esc(flabel)}</h3><{tag}>" + "".join(f"<li>{esc(v)}</li>" for v in val) + f"</{tag}>")

        for field, flabel in [("government_scheme_alignment", "Government Scheme Alignment"),
                               ("indian_legal_references", "Indian Legal References"),
                               ("indian_case_references", "Indian Case References"),
                               ("references", "References"),
                               ("keywords", "Keywords")]:
            val = report.get(field)
            if val:
                if isinstance(val, list):
                    sections.append(f"<p><strong>{esc(flabel)}:</strong> {esc(', '.join(val))}</p>")
                else:
                    sections.append(f"<p><strong>{esc(flabel)}:</strong> {esc(str(val))}</p>")

    body = "\n".join(sections)
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{esc(title)}</title>
<style>
body {{ font-family: 'Times New Roman', serif; margin: 1in; font-size: 11pt; line-height: 1.6; }}
h1 {{ font-size: 20pt; text-align: center; margin-bottom: 10px; }}
h2 {{ font-size: 14pt; margin-top: 25px; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
h3 {{ font-size: 12pt; margin-top: 18px; }}
h4 {{ font-size: 11pt; margin-top: 12px; }}
p {{ text-align: justify; margin: 4px 0; }}
ul, ol {{ margin-left: 20px; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 20px 0; }}
table {{ page-break-inside: avoid; }}
</style></head><body>
{body}
</body></html>"""


def _table(headers: list, rows: list) -> str:
    hdr = "<tr>" + "".join(f"<th>{escape(str(h))}</th>" for h in headers) + "</tr>"
    body = ""
    for row in rows:
        body += "<tr>" + "".join(f"<td>{r}</td>" for r in row) + "</tr>"
    return f'<table border="1" cellpadding="5" cellspacing="0" style="border-collapse:collapse; width:100%; font-size:10pt;"><thead>{hdr}</thead><tbody>{body}</tbody></table>'


def _fallback_docx(report: dict) -> bytes:
    text = f"""
{report.get('title', 'EIPR Report')}

Abstract:
{report.get('abstract', '')}

Case Study:
{json.dumps(report.get('case_study', {}), indent=2)}

Keywords: {', '.join(report.get('keywords', []))}
"""
    return text.encode("utf-8")
